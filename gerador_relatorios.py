import os
from fpdf import FPDF
from docx import Document
import pandas as pd
from datetime import datetime

# Criar pasta de saída
OUTPUT_DIR = 'relatorios'
os.makedirs(OUTPUT_DIR, exist_ok=True)

def gerar_nome_arquivo(extensao):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(OUTPUT_DIR, f"relatorio_{timestamp}.{extensao}")

def gerar_pdf(dados):
    caminho = gerar_nome_arquivo("pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Relatório SENKAS", ln=True, align='C')

    for chave, valor in dados.items():
        pdf.cell(200, 10, txt=f"{chave}: {valor}", ln=True)

    pdf.output(caminho)
    return caminho

def gerar_word(dados):
    caminho = gerar_nome_arquivo("docx")
    doc = Document()
    doc.add_heading('Relatório SENKAS', 0)

    for chave, valor in dados.items():
        doc.add_paragraph(f"{chave}: {valor}")

    doc.save(caminho)
    return caminho

def gerar_excel(dados):
    caminho = gerar_nome_arquivo("xlsx")
    df = pd.DataFrame([dados])
    df.to_excel(caminho, index=False)
    return caminho
